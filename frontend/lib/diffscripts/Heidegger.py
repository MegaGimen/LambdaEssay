import argparse
import sys
import os
from flask import Flask, request, jsonify
from docx import Document

app = Flask(__name__)

def get_run_style(run):
    color = None
    try:
        if run.font.color and run.font.color.rgb:
            color = str(run.font.color.rgb)
    except:
        pass

    return {
        'text': run.text,
        'bold': run.bold,
        'italic': run.italic,
        'underline': run.underline,
        'font_name': run.font.name,
        'font_size': run.font.size.pt if run.font.size else None,
        'color': color,
    }

def get_para_style(para):
    return {
        'alignment': str(para.alignment),
        'style': para.style.name,
        'runs': normalize_runs([get_run_style(run) for run in para.runs])
    }

def normalize_runs(runs):
    if not runs:
        return []
    
    normalized = []
    current_run = None
    
    for run in runs:
        if current_run is None:
            current_run = run
        else:
            # Check if style matches (excluding text)
            style_match = True
            for key in run:
                if key != 'text' and run[key] != current_run[key]:
                    style_match = False
                    break
            
            if style_match:
                current_run['text'] += run['text']
            else:
                normalized.append(current_run)
                current_run = run
    
    if current_run:
        normalized.append(current_run)
        
    # Filter out empty runs that might have been created or existed
    return [r for r in normalized if r['text']]

def extract_content(docx_path_or_stream):
    doc = Document(docx_path_or_stream)
    content = []
    for para in doc.paragraphs:
        content.append(get_para_style(para))
    
    # Also checking tables
    for table in doc.tables:
        table_content = []
        for row in table.rows:
            row_content = []
            for cell in row.cells:
                cell_content = []
                for para in cell.paragraphs:
                    cell_content.append(get_para_style(para))
                row_content.append(cell_content)
            table_content.append(row_content)
        content.append({'type': 'table', 'data': table_content})
        
    return content

def compare_docs(doc1_source, doc2_source):
    try:
        content1 = extract_content(doc1_source)
        content2 = extract_content(doc2_source)
        print(content1)
        print(content2)

        print(content1==content2)
        
        if content1 == content2:
            return True, "No differences found."
        else:
            return False, "Differences found in content or style."
    except Exception as e:
        return False, f"Error processing files: {str(e)}"

@app.route('/compare', methods=['POST'])
def compare_endpoint():
    if 'file1' not in request.files or 'file2' not in request.files:
        return jsonify({"error": "Please upload both file1 and file2"}), 400
    
    file1 = request.files['file1']
    file2 = request.files['file2']
    
    same, message = compare_docs(file1, file2)
    return jsonify({"identical": same, "message": message})

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description='Compare two docx files for content and style differences.')
    parser.add_argument('--docx1', help='Path to the first docx file')
    parser.add_argument('--docx2', help='Path to the second docx file')
    parser.add_argument('--port', type=int, default=5000, help='Port to run the server on')
    
    args = parser.parse_args()
    
    if args.docx1 and args.docx2:
        if not os.path.exists(args.docx1) or not os.path.exists(args.docx2):
            print("Error: One or both file paths do not exist.")
            sys.exit(1)
            
        same, message = compare_docs(args.docx1, args.docx2)
        print(f"Identical: {same}")
        print(f"Message: {message}")
    else:
        print(f"Starting server on port {args.port}...")
        app.run(host='0.0.0.0', port=args.port, debug=False)
